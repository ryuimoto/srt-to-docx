import os
import re
import sys
from docx import Document

# 自然な並び順でソートするためのキー関数
def natural_sort_key(filename):
    name, ext = os.path.splitext(filename)
    return [int(tok) if tok.isdigit() else tok for tok in re.split(r'(\d+)', name)]

def extract_text_from_srt(srt_file):
    with open(srt_file, "r", encoding="utf-8") as file:
        content = file.read()

    entries = re.findall(r'\d+\n(\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3})\n(.+?)(?=\n\d+|$)', content, re.DOTALL)

    with_timecodes = ""
    without_timecodes = ""

    for timestamp, text in entries:
        text = re.sub(r'\n+', '\n', text).strip()
        with_timecodes += f"{timestamp}\n{text}\n\n"
        without_timecodes += text + "\n"

    return with_timecodes.strip(), without_timecodes.strip()

def create_docx_for_folder(srt_folder):
    try:
        files = os.listdir(srt_folder)
        mp4_files = sorted([f for f in files if f.endswith('.mp4')], key=natural_sort_key)
        srt_files = sorted([f for f in files if f.endswith('.srt')], key=natural_sort_key)
    except Exception as e:
        print(f"フォルダの読み込みに失敗: {srt_folder}, エラー: {e}")
        return

    if not mp4_files:
        print(f"MP4ファイルが見つかりません: {srt_folder}")
        return

    mp4_filename = os.path.splitext(mp4_files[0])[0]

    docx_with_time = Document()
    docx_without_time = Document()

    all_with_timecodes = ""
    all_without_timecodes = ""

    for srt_file in srt_files:
        srt_path = os.path.join(srt_folder, srt_file)
        with_time, without_time = extract_text_from_srt(srt_path)
        all_with_timecodes += with_time + "\n\n"
        all_without_timecodes += without_time + "\n"
        
        # SRTファイルを削除
        os.remove(srt_path)
        print(f"削除完了: {srt_path}")

    docx_with_time.add_paragraph(all_with_timecodes.strip())
    docx_without_time.add_paragraph(all_without_timecodes.strip())

    with_time_filename = os.path.join(srt_folder, f"{mp4_filename}.docx")
    without_time_filename = os.path.join(srt_folder, f"{mp4_filename} ※タイムコードなし.docx")

    docx_with_time.save(with_time_filename)
    docx_without_time.save(without_time_filename)

    print(f"保存完了: {with_time_filename}")
    print(f"保存完了: {without_time_filename}")

def process_multiple_folders(base_folder):
    for root, dirs, files in os.walk(base_folder):
        if any(f.endswith('.srt') for f in files):
            create_docx_for_folder(root)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python3 srt_to_docx.py <フォルダパス>")
        sys.exit(1)
    
    base_folder_path = sys.argv[1]
    process_multiple_folders(base_folder_path)